"""
Resume skills extractor — parses PDF / DOCX / TXT resumes and infers a
canonical list of skills via the Azure AI Foundry Grok endpoint.

Pipeline:
  1. Extract plain text from the uploaded resume (PDF, DOCX, or TXT)
  2. Send text to the Grok model with a structured extraction prompt
  3. Parse the JSON response and return a deduplicated list of skills

Exposed:
  extract_text_from_resume(filename, data) -> str
  infer_skills_from_resume_text(text) -> list[str]
  process_resume_upload(filename, data) -> list[str]        # one-shot convenience
"""

from __future__ import annotations

import io
import json
import logging
import re
from typing import List

import httpx

from ..config import settings

logger = logging.getLogger(__name__)

AI_MODEL = settings.azure_ai_model or "grok-4-fast-reasoning"

# Cap the text we send to the model — resumes this size are already rare,
# and we want to stay well under any per-request token limits.
MAX_RESUME_CHARS = 24_000


# ---------------------------------------------------------------------------
# Text extraction (PDF / DOCX / TXT)
# ---------------------------------------------------------------------------

def _extract_pdf(data: bytes) -> str:
    """Extract PDF text using PyMuPDF (fitz) with Tesseract OCR fallback.

    Strategy:
      1. Try `page.get_text("text", sort=True)` — fast path for text-layer PDFs.
         PyMuPDF handles multi-column layouts and tables well here.
      2. If the combined text is too sparse (almost certainly a scanned /
         image-only PDF), render each page at 200 DPI and run Tesseract OCR.

    The fast path is instantaneous; OCR takes a few seconds per page but is
    only invoked when the text path has nothing to return.
    """
    import pymupdf  # deferred import

    pages: List[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            try:
                text = page.get_text("text", sort=True) or ""
            except Exception as e:  # noqa: BLE001
                logger.warning("pymupdf get_text failed on page %d: %s", page_num, e)
                text = ""
            if text.strip():
                pages.append(text)

    combined = "\n\n".join(pages).strip()

    # If the text layer was effectively empty or very sparse, fall back to OCR.
    # Threshold of 150 chars catches image-only PDFs and "PDFs" that are really
    # just a wrapper around a scanned image.
    if len(combined) < 150:
        logger.info(
            "PDF text layer sparse (%d chars). Falling back to Tesseract OCR.",
            len(combined),
        )
        try:
            combined = _ocr_pdf(data)
        except Exception as e:  # noqa: BLE001
            logger.warning("OCR fallback failed: %s", e)

    return combined


def _ocr_pdf(data: bytes) -> str:
    """Render each PDF page to an image and OCR it with Tesseract.

    Used when PyMuPDF's text layer is empty (scanned / image-only PDFs).
    """
    import pymupdf
    import pytesseract
    from PIL import Image

    pages: List[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            try:
                # 200 DPI gives a good quality / speed trade-off for OCR.
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img) or ""
            except Exception as e:  # noqa: BLE001
                logger.warning("Tesseract failed on page %d: %s", page_num, e)
                text = ""
            if text.strip():
                pages.append(text)
    out = "\n\n".join(pages).strip()
    logger.info("OCR fallback extracted %d chars across %d pages", len(out), len(pages))
    return out


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(data))
    parts: List[str] = [p.text for p in doc.paragraphs if p.text]
    # Include table cells too — many resumes use tables for layout
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    # Be generous with decoding — resumes can be exported as latin-1 too
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text_from_resume(filename: str, data: bytes) -> str:
    """Extract plain text from a resume file (PDF, DOCX, or TXT)."""
    fname = (filename or "").lower()
    if fname.endswith(".pdf"):
        text = _extract_pdf(data)
    elif fname.endswith(".docx"):
        text = _extract_docx(data)
    elif fname.endswith(".txt"):
        text = _extract_txt(data)
    else:
        # Fall back on the bytes signature — magic numbers
        if data.startswith(b"%PDF-"):
            text = _extract_pdf(data)
        elif data.startswith(b"PK"):  # .docx is a zip
            try:
                text = _extract_docx(data)
            except Exception:
                raise ValueError(
                    "Unsupported file type. Please upload a PDF, DOCX, or TXT resume."
                )
        else:
            text = _extract_txt(data)

    # Normalize whitespace — the AI model does NOT care about formatting
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    if len(text) > MAX_RESUME_CHARS:
        text = text[:MAX_RESUME_CHARS]

    return text


# ---------------------------------------------------------------------------
# Grok AI skill inference
# ---------------------------------------------------------------------------

SKILL_EXTRACTION_SYSTEM = (
    "You are a world-class senior engineering recruiter with encyclopedic, "
    "up-to-date knowledge of every modern tech stack, cloud service, "
    "framework, database, data/ML tool, and engineering practice. Your job "
    "is to read a developer's resume and build a rigorous, complete picture "
    "of their skills.\n\n"
    "Two modes, both mandatory:\n"
    "  1. DIRECT: when skills are named explicitly (Skills section, tool list, "
    "     explicit mentions) — include them.\n"
    "  2. INFERRED: when the resume describes work done, REASON about what "
    "     technologies/tools/practices the engineer must have used to deliver "
    "     that outcome, and include those too. Most senior resumes under-list "
    "     skills — the real evidence is in project descriptions.\n\n"
    "For every project / job / internship bullet, think step by step: what "
    "systems would this work require? What languages, frameworks, databases, "
    "cloud services, infra patterns? Include inferred skills with the same "
    "confidence as directly listed ones, provided the evidence is specific "
    "enough (quantified scale, named companies, concrete outcomes, named "
    "subsystems). Skip inferences that would be pure speculation.\n\n"
    "If the resume is clearly non-technical (legal, medical, marketing, etc.), "
    "still do your job competently and extract skills appropriate to that "
    "field — but your primary expertise is software engineering and that's "
    "where your inference should shine."
)

SKILL_EXTRACTION_PROMPT = """\
Read the resume carefully. Scan every section: Skills, Experience, Projects,
Internships, Education, Publications, Certifications, Awards.

Your output: a JSON object {"skills": [...]} — an array of concise skill tags
(1–4 words each). No prose, no markdown, JSON only.

================================================================================
HOW TO EXTRACT — BOTH PATHS
================================================================================

PATH 1 — Direct extraction
  If a skill is named (in Skills list, bullet text, or a project tech stack),
  include it. Example: "Built an app with React, Node.js, PostgreSQL" →
  React, Node.js, PostgreSQL.

PATH 2 — Inference from work described
  For each project/role bullet, reason: "To have shipped this outcome, the
  engineer would have needed <technologies>." Include the inferred skills.

  Inference examples (study the pattern):

  Bullet: "Built a real-time multiplayer matchmaking service handling 50k
          concurrent players across AWS us-east and eu-west."
  Infer: WebSockets, Redis (pub/sub), Distributed Systems, AWS, ElastiCache,
         Multi-Region Deployment, Load Balancing, Horizontal Scaling, Latency
         Optimization, Go or Node.js (pick based on context).

  Bullet: "Reduced p99 API latency from 800ms to 120ms by rewriting hot path
          in Go and adding Redis cache layer."
  Infer: Go, Redis, Performance Tuning, Caching Strategy, Profiling,
         API Optimization, p99 Latency.

  Bullet: "Migrated monolithic Rails app to 12 microservices on Kubernetes
          with event-driven communication via Kafka."
  Infer: Ruby on Rails, Microservices, Kubernetes, Kafka, Event-Driven
         Architecture, Service Mesh, Containerization, Distributed Tracing,
         System Design.

  Bullet: "Led the data pipeline that ingests 2B rows/day from Kafka into
          Snowflake and powers the ML feature store."
  Infer: Kafka, Snowflake, Data Pipelines, ETL, Batch Processing, ML Feature
         Engineering, Big Data, SQL, Airflow (likely).

  Bullet: "Interned at Hyperverge AI team, working on KYC document labeling
          and model contracts."
  Infer: Computer Vision (KYC is vision/OCR), Machine Learning, Model
         Deployment, Data Labeling, MLOps (lightly).

  Bullet: "Published paper on Transformer variants for low-resource NMT."
  Infer: Deep Learning, Transformers, NLP, PyTorch (likely), Research,
         Sequence Models, Neural Machine Translation.

  Bullet: "Shipped iOS app with 4.7★ rating and 200k MAUs."
  Infer: iOS Development, Swift (or SwiftUI), Mobile Development, App Store
         Deployment, Product Analytics.

================================================================================
CANONICALIZATION
================================================================================

Always use the canonical form. Deduplicate aggressively.

  react / ReactJS                → React
  nodejs / node                  → Node.js
  js / Javascript                → JavaScript
  ts                             → TypeScript
  postgres / psql / PG           → PostgreSQL
  k8s                            → Kubernetes
  tf (in ML context)             → TensorFlow
  ci/cd / ci-cd                  → CI/CD
  ms sql / mssql                 → SQL Server
  ipr / ip rights / IP Law       → Intellectual Property Law  (pick one)
  aws lambda / lambda            → AWS Lambda
  gcp                            → Google Cloud

Never list both "IPR" and "Intellectual Property Law" — pick the canonical one.

================================================================================
WHAT TO EXCLUDE
================================================================================

- Generic soft skills: teamwork, communication, leadership, problem solving
  (unless uniquely demonstrated, like "Led 12-person engineering team" →
  "Engineering Leadership").
- Single vague words: "AI", "software", "development", "technology" — unless
  specific ("LLM Fine-Tuning", "Mobile Development" are fine).
- Pure buzzwords without any supporting context.
- Your own uncertainty — if you aren't confident about an inference, omit it.

================================================================================
OUTPUT
================================================================================

- 1–4 word tags, title-cased properly (React, not react; PostgreSQL, not postgres).
- Ordered by centrality: most-evidenced and most-recent first.
- Up to 30 skills. Quality over quantity.
- JSON only. No explanations. No markdown fences.

Example output:
{"skills": ["Python", "React", "TypeScript", "PostgreSQL", "AWS", "Docker", "Kubernetes", "Microservices", "Kafka", "Redis", "CI/CD", "GraphQL", "System Design"]}

Resume text:
---
%s
---
"""


def _parse_skills_from_response(raw: str) -> List[str]:
    """Parse the model's JSON response into a clean skill list."""
    text = raw.strip()
    # Strip markdown code fences if the model added them
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```\s*$", "", text)
    try:
        obj = json.loads(text)
    except json.JSONDecodeError:
        # Try to pull a JSON object substring out
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            raise ValueError("AI response did not contain valid JSON")
        obj = json.loads(match.group(0))

    raw_skills = obj.get("skills", []) if isinstance(obj, dict) else []
    if not isinstance(raw_skills, list):
        raise ValueError("AI response 'skills' field was not a list")

    # Clean + dedupe preserving order
    seen: set[str] = set()
    cleaned: List[str] = []
    for s in raw_skills:
        if not isinstance(s, str):
            continue
        s2 = s.strip()
        if not s2:
            continue
        key = s2.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(s2)
    return cleaned[:30]


async def infer_skills_from_resume_text(text: str) -> List[str]:
    """Call Grok on Azure AI Foundry and return the inferred skills."""
    if not text.strip():
        return []

    if not settings.azure_ai_api_key or not settings.azure_ai_endpoint:
        logger.warning("Azure AI not configured — cannot infer skills")
        raise RuntimeError(
            "AI service is not configured. Please contact your administrator."
        )

    prompt = SKILL_EXTRACTION_PROMPT % text
    payload = {
        "model": AI_MODEL,
        "messages": [
            {"role": "system", "content": SKILL_EXTRACTION_SYSTEM},
            {"role": "user", "content": prompt},
        ],
        # Grok-4-fast is a reasoning model — it uses tokens for internal
        # chain-of-thought BEFORE producing the JSON output. Give it room.
        "max_tokens": 4000,
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
    }
    headers = {
        "Content-Type": "application/json",
        "api-key": settings.azure_ai_api_key,
    }

    logger.info("Calling Azure AI Foundry (%s) for resume skill extraction", AI_MODEL)
    async with httpx.AsyncClient(timeout=httpx.Timeout(120.0)) as client:
        resp = await client.post(settings.azure_ai_endpoint, headers=headers, json=payload)
        resp.raise_for_status()

    data = resp.json()
    raw = data["choices"][0]["message"]["content"]
    skills = _parse_skills_from_response(raw)
    logger.info("Extracted %d skills from resume", len(skills))
    return skills


async def process_resume_upload(filename: str, data: bytes) -> List[str]:
    """One-shot helper: extract text + infer skills.

    Logs the extracted text length so we can sanity-check extraction quality
    when a resume returns few/no skills. Very short text usually means the PDF
    is scanned (image-only) or otherwise unreadable.
    """
    text = extract_text_from_resume(filename, data)
    logger.info(
        "Resume extraction: filename=%s file_bytes=%d text_chars=%d",
        filename, len(data), len(text),
    )
    if not text:
        raise ValueError(
            "We couldn't read any text from this file. If it's a scanned PDF, "
            "please upload a text-based PDF or a DOCX instead."
        )
    if len(text) < 200:
        # Extremely short extraction — likely a scanned PDF with just a name on top
        raise ValueError(
            "This file looks almost empty. If it's a scanned or image-based PDF, "
            "please export it as a text PDF or DOCX and try again."
        )
    return await infer_skills_from_resume_text(text)
